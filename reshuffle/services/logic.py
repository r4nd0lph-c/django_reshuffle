import os
import shutil

from random import shuffle, choice
from math import ceil, sqrt
from datetime import datetime

from django_reshuffle.settings import MEDIA_ROOT
from reshuffle.models import *
from reshuffle.services import unique_key

import docx
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from lxml import etree
import latex2mathml.converter

DOCS_ROOT = os.path.join(MEDIA_ROOT, 'docs')
MML2OMML_ROOT = os.path.join(MEDIA_ROOT, 'MML2OMML')

ARCHIVE_NAME = 'reshuffle_tasks'
FOLDER_NAME_Q = 'задания'
FOLDER_NAME_A = 'ответы'
ARCHIVE_LIFETIME = 24  # hours

NULL_ITEM = None


def get_time_create(obj):
    """ FUNCTION that return file time create"""
    return datetime.fromtimestamp(os.path.getmtime(obj))


def clean_trash(username):
    """
    FUNCTION to remove old user's archive
    > archive is old if difference between archive time_create and time_now is bigger than ARCHIVE_LIFETIME

    INPUT: string username - user's name for directory

    RESULT: remove directory with old archive
    OUTPUT: boolean is_cleaning
    """

    archive = None
    if os.path.exists(os.path.join(DOCS_ROOT, username, ARCHIVE_NAME + '.zip')):
        archive = os.path.join(DOCS_ROOT, username, ARCHIVE_NAME + '.zip')  # find old archive if it exists

    if archive is not None:
        time_create = get_time_create(archive)
        time_exist = abs(datetime.now() - time_create).total_seconds() / 3600  # time exist in hours

        if time_exist > ARCHIVE_LIFETIME:  # THE TIME HAS COME AND SO HAVE I...
            shutil.rmtree(os.path.join(DOCS_ROOT, username))  # remove  directory with archive
            return True
        else:
            return False
    else:
        return True


def latex_to_word(latex_input):
    """ FUNCTION that transform LaTeX input to MathML output """
    mathml = latex2mathml.converter.convert(latex_input)
    tree = etree.fromstring(mathml)
    xslt = etree.parse(os.path.join(MML2OMML_ROOT, 'MML2OMML.XSL'))
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def docx_settings(document):
    """
    FUNCTION to pre-configure document formatting

    INPUT: 'document' is a docx.Document() for formatting

    RESULT: custom sections, font
    OUTPUT: head object
    """

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    head = document.add_paragraph()
    head.style = document.styles['Normal']
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return head


def insert_hr(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'gray')
    pBdr.append(bottom)


def create_docx_question(subject, data, create_date):
    """
    FUNCTION for creating one .docx question-file

    INPUT: 'subject' is a (class Subject) object from reshuffle/models.py
    INPUT: dict 'data' is a result of variant_data() function
    INPUT: 'create_date' is a string with format: '%Y_%m_%d_%H_%M_%S_%f'

    RESULT: + one .docx question-file in directory
    """

    # create document from template
    document = docx.Document(os.path.join(MEDIA_ROOT, 'template.docx'))

    # personal num & cutting part
    document._body.clear_content()

    personal_num = document.add_paragraph()
    personal_num.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    personal_num.alignment = 1

    personal_num.add_run('Номер личного дела: ').bold = True
    personal_num.add_run('__ ' * 9)

    cutting_part = document.add_paragraph()
    cutting_part.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    cutting_part.add_run('Дата: ').bold = True
    date_list = create_date.split('_')
    cutting_part.add_run('__.{}.{}'.format(date_list[1], date_list[0]))

    cutting_part.add_run('     Код: ').bold = True
    cutting_part.add_run(data['unique_key'])

    cutting_part.add_run('     ФИО:').bold = True
    cutting_part.add_run(' ' * 60 + 'Подпись:').bold = True

    insert_hr(cutting_part)
    # -----

    # head university & test info
    head = document.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    head.add_run('\nМИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ\n').bold = True
    head.add_run('Федеральное государственное автономное образовательное учреждение\nвысшего образования\n')
    head.add_run('«МОСКОВСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»\n\n')

    test_info = head.add_run('Тест по ' + subject.case_dative + '\n' + 'Вариант № ' + data['unique_key'] + '\n\n\n')
    test_info.bold = True
    test_info.font.size = Pt(13)
    # -----

    # code in header
    for section in document.sections:
        header = section.header
        header.paragraphs[0].alignment = 1
        header.paragraphs[0].text = data['unique_key']
    # -----

    # general info add
    general_info = subject.header
    if general_info:
        head.add_run(general_info['title'] + '\n').bold = True
        head.add_run(general_info['text'] + '\n\n').italic = True
    # -----

    # FILL QUESTION-DOC.docx WITH INFO FUNCTION
    n_len = len(str(len(data['body'])))
    num = 0

    # prepare to gen char-num
    def add_part(p_i, j):
        new_part = document.add_paragraph()
        new_part.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_part.add_run('Часть ' + p_i[j][0] + '\n').bold = True
        new_part.add_run(p_i[j][3] + '\n').italic = True

    accum = 0
    jumper = 0
    parts_info = []
    parts = subject.parts
    if parts:
        max_k = -1
        for key in parts.keys():
            if parts[key]['number'] > max_k:
                max_k = parts[key]['number']
            parts_info.append((key, parts[key]['number'], parts[key]['number'] + accum, parts[key]['help_text']))
            accum += parts[key]['number']
        n_len = len(str(max_k))
        add_part(parts_info, jumper)
    # -----

    for item in data['body']:
        num += 1
        # gen char-num
        char_num = str(num).zfill(n_len)
        if parts:
            if not (parts_info[jumper][2] - parts_info[jumper][1] <= num <= parts_info[jumper][2]):
                jumper += 1
                add_part(parts_info, jumper)
            char_num = '{}{}'.format(parts_info[jumper][0],
                                     str(num + parts_info[jumper][1] - parts_info[jumper][2]).zfill(n_len))
        # -----
        if item != NULL_ITEM:
            task = document.add_paragraph()
            n_text = task.add_run('[' + char_num + '] ')
            n_text.bold = True
            task_text = task.add_run(item['task'].text + '\n')
            n_text.font.size, task_text.font.size = Pt(12), Pt(12)

            if item['task'].latex != '':
                task.add_run(' ')
                word_math = latex_to_word(item['task'].latex)
                task._element.append(word_math)
                task.add_run('\n')

            if item['task'].image != '':
                r = task.add_run()
                r.add_picture(item['task'].image, width=Inches(6))
                task.add_run('\n')

            if len(item['options']) != 1:
                i = 0
                for opt in item['options']:
                    i += 1
                    task.add_run('  {}) {}'.format(i, opt.text))
                    if opt.latex != '':
                        task.add_run(' ')
                        word_math = latex_to_word(opt.latex)
                        task._element.append(word_math)
                    if opt.image != '':
                        r = task.add_run()
                        r.add_picture(opt.image, width=Inches(4))
                    task.add_run('\n')
    # ----- -----

    # bottom part
    insert_hr(document.paragraphs[-1])

    bottom_part = document.add_paragraph()
    bottom_part.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom_part.add_run('\nОтмена ошибочных меток').bold = True

    change_table = document.add_table(rows=1, cols=2)
    table_heading = change_table.rows[0].cells
    table_heading[0].paragraphs[0].text = 'Номер задания'
    table_heading[1].paragraphs[0].text = 'Исправленный ответ'
    for i in range(0, 5):
        row = change_table.add_row().cells
        row[0].text = '№ ____'
        row[1].text = 'Отв. ____'
    change_table.style = 'change_table_style'

    exam_check = document.add_paragraph()
    exam_check.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_check.add_run('\n\nПроверка выполненной работы').bold = True

    exam_table = document.add_table(rows=1, cols=2)
    table_heading = exam_table.rows[0].cells
    table_heading[0].paragraphs[0].text = 'Количество решенных задач:'
    table_heading[1].paragraphs[0].text = 'Подпись экзаменатора:'
    exam_table.style = 'exam_table_style'
    # -----

    # save document
    document.save(
        '{}\\{}\\{}\\{}_{}_задания.docx'.format(DOCS_ROOT, create_date, FOLDER_NAME_Q,
                                                subject.case_nominative, data['unique_key']))
    # -----


def create_docx_answer(subject, data, create_date):
    """
    FUNCTION for creating one .docx answer-file

    INPUT: 'subject' is a (class Subject) object from reshuffle/models.py
    INPUT: dict 'data' is a result of variant_data() function
    INPUT: 'create_date' is a string with format: '%Y_%m_%d_%H_%M_%S_%f'

    RESULT: + one .docx answer-file in directory
    """

    document = docx.Document()
    head = docx_settings(document)

    test_form = head.add_run('Ключи ответов тестов по ' + subject.case_dative + '\n')
    test_form.bold = True
    test_form.font.size = Pt(13)

    space = head.add_run('\n')
    space.font.size = Pt(13)

    variant_num = head.add_run('Вариант № ' + data['unique_key'] + '\n')
    variant_num.bold = True
    variant_num.font.size = Pt(13)

    # FILL ANSWER-DOC.docx WITH INFO FUNCTION
    n = len(data['body'])
    n_len = len(str(n))

    cols_count = ceil(sqrt(n))
    while n % cols_count != 0:
        cols_count -= 1

    rows_count = n // cols_count * 2

    # prepare to gen char-num
    accum = 0
    jumper = 0
    parts_info = []
    parts = subject.parts
    if parts:
        max_k = -1
        for key in parts.keys():
            if parts[key]['number'] > max_k:
                max_k = parts[key]['number']
            parts_info.append((key, parts[key]['number'], parts[key]['number'] + accum, parts[key]['help_text']))
            accum += parts[key]['number']
        n_len = len(str(max_k))
    # -------------------------
    # Table data in a form of list
    data_table = []
    for i in range(0, rows_count // 2):
        row = []
        for j in range(0, cols_count):
            # gen char-num
            num = i * cols_count + j + 1
            char_num = str(num).zfill(n_len)
            if parts:
                if not (parts_info[jumper][2] - parts_info[jumper][1] <= num <= parts_info[jumper][2]):
                    jumper += 1
                char_num = '{}{}'.format(parts_info[jumper][0],
                                         str(num + parts_info[jumper][1] - parts_info[jumper][2]).zfill(n_len))
            # -------------------------
            row.append('Задание ' + char_num)
        data_table.append(row)

        row = []
        for j in range(0, cols_count):
            if data['body'][i * cols_count + j] != NULL_ITEM:
                opt_collection = data['body'][i * cols_count + j]['options']
                if len(opt_collection) != 1:
                    otv_n = 0
                    for obj in opt_collection:
                        otv_n += 1
                        if obj.is_answer:
                            break
                    if otv_n == 0:
                        row.append('Письменный ответ')
                    else:
                        row.append(otv_n)
                else:
                    row.append('{} {}'.format(opt_collection[0].text, opt_collection[0].latex))
            else:
                row.append('NONE')
        data_table.append(row)

    # Creating a table object
    table = document.add_table(rows=rows_count, cols=cols_count)

    # # Adding data from the list to the table
    j = 0
    for data_row in data_table:
        row = table.rows[j].cells
        for i in range(0, cols_count):
            # p = row[i].add_paragraph(str(data_row[i]))
            p = row[i].paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.text = str(data_row[i])
            if j % 2 == 0:
                run = p.runs[0]
                run.font.bold = True
        j += 1

    # Adding style to a table
    table.style = 'Table Grid'
    # --------------------------------------------------

    document.save(
        '{}\\{}\\{}\\{}_{}_ответы.docx'.format(DOCS_ROOT, create_date, FOLDER_NAME_A,
                                               subject.case_nominative, data['unique_key']))


def variant_data(tasks_queryset, n):
    """
        FUNCTION for selecting info from the database

        INPUT: queryset 'tasks_queryset' with all tasks for subject
        INPUT: int 'n' - count of tasks in subject

        OUTPUT: dict 'data' with 'unique_key' and 'body'
        > 'unique_key' is a string key
        > 'body' is a list of dicts with 'task' and 'options'
        > > 'task' is a (class Tasks) object from reshuffle/models.py
        > > 'options' is a shuffle list with (class Options) objects from reshuffle/models.py for this 'task'
    """

    data = {'unique_key': unique_key.create(), 'body': []}

    for i in range(0, n):
        pool_tasks_i = tasks_queryset.filter(num=(i + 1))
        if pool_tasks_i:
            task_i = choice(pool_tasks_i)
            pool_options_i = Options.objects.filter(task_fk=task_i.id)

            options_i = []
            for item in pool_options_i:
                options_i.append(item)
            shuffle(options_i)

            data['body'].append({'task': task_i, 'options': options_i})
        else:
            data['body'].append(NULL_ITEM)

    return data


def main(cleaned_data, username):
    """
    MAIN FUNCTION

    INPUT: cleaned_data from CreationForm (render in creation.html)
    INPUT: string username - user's name for directory

    RESULT: create new archives with .docx files
    """

    create_date = datetime.now().strftime('%Y_%m_%d_%H_%M_%S_%f')  # future name of tmp directory

    os.mkdir(os.path.join(DOCS_ROOT, create_date))  # create new tmp directory
    os.mkdir(os.path.join(DOCS_ROOT, create_date, FOLDER_NAME_Q))  # nested directories
    os.mkdir(os.path.join(DOCS_ROOT, create_date, FOLDER_NAME_A))  # nested directories

    if not os.path.exists(os.path.join(DOCS_ROOT, username)):
        os.mkdir(os.path.join(DOCS_ROOT, username))  # create if not exist user's directory for archive

    subject = cleaned_data['subject']
    amount = cleaned_data['amount']

    tasks_queryset = Tasks.objects.filter(subject_fk=subject.id)

    # create .docx files in tmp directory
    for i in range(0, amount):
        data = variant_data(tasks_queryset, subject.tasks_number)
        create_docx_question(subject, data, create_date)
        create_docx_answer(subject, data, create_date)

    shutil.make_archive(os.path.join(DOCS_ROOT, username, ARCHIVE_NAME),
                        'zip',
                        os.path.join(DOCS_ROOT, create_date))  # create new archive

    shutil.rmtree(os.path.join(DOCS_ROOT, create_date))  # remove tmp directory with .docx files


if __name__ == '__main__':
    pass
